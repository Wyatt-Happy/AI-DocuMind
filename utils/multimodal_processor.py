import os
import json
import base64
from PIL import Image
import io
import docx
import fitz

class MultimodalProcessor:
    def __init__(self, output_dir="data/multimodal"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(os.path.join(output_dir, "images"), exist_ok=True)
        os.makedirs(os.path.join(output_dir, "tables"), exist_ok=True)
    
    def extract_images_from_word(self, file_path):
        """从Word文件中提取图片"""
        try:
            doc = docx.Document(file_path)
            images = []
            
            # 提取内联形状中的图片
            for i, shape in enumerate(doc.inline_shapes):
                if shape.type == docx.enum.text.WD_INLINE_SHAPE.PICTURE:
                    # 获取图片数据
                    image_data = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                    image_part = doc.part.related_parts[image_data]
                    image_bytes = image_part._blob
                    
                    # 保存图片
                    image_filename = f"word_image_{i+1}.png"
                    image_path = os.path.join(self.output_dir, "images", image_filename)
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)
                    
                    images.append({
                        'filename': image_filename,
                        'path': image_path,
                        'type': 'inline',
                        'index': i+1
                    })
            
            return images
        except Exception as e:
            print(f"从Word文件提取图片失败: {e}")
            return []
    
    def extract_images_from_pdf(self, file_path):
        """从PDF文件中提取图片"""
        try:
            doc = fitz.open(file_path)
            images = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images(full=True)
                
                for i, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # 保存图片
                    image_filename = f"pdf_image_page{page_num+1}_{i+1}.png"
                    image_path = os.path.join(self.output_dir, "images", image_filename)
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)
                    
                    images.append({
                        'filename': image_filename,
                        'path': image_path,
                        'type': 'pdf',
                        'page': page_num+1,
                        'index': i+1
                    })
            
            return images
        except Exception as e:
            print(f"从PDF文件提取图片失败: {e}")
            return []
    
    def extract_tables_from_word(self, file_path):
        """从Word文件中提取表格"""
        try:
            doc = docx.Document(file_path)
            tables = []
            
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                # 保存表格数据
                table_filename = f"word_table_{i+1}.json"
                table_path = os.path.join(self.output_dir, "tables", table_filename)
                with open(table_path, "w", encoding="utf-8") as f:
                    json.dump(table_data, f, ensure_ascii=False, indent=2)
                
                tables.append({
                    'filename': table_filename,
                    'path': table_path,
                    'data': table_data,
                    'index': i+1
                })
            
            return tables
        except Exception as e:
            print(f"从Word文件提取表格失败: {e}")
            return []
    
    def extract_tables_from_pdf(self, file_path):
        """从PDF文件中提取表格"""
        try:
            doc = fitz.open(file_path)
            tables = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                # 使用页面文本提取表格（简单方法）
                text = page.get_text()
                lines = text.split('\n')
                
                # 简单检测表格（基于行对齐）
                table_data = []
                in_table = False
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        if in_table and table_data:
                            # 保存表格
                            table_filename = f"pdf_table_page{page_num+1}_{len(tables)+1}.json"
                            table_path = os.path.join(self.output_dir, "tables", table_filename)
                            with open(table_path, "w", encoding="utf-8") as f:
                                json.dump(table_data, f, ensure_ascii=False, indent=2)
                            
                            tables.append({
                                'filename': table_filename,
                                'path': table_path,
                                'data': table_data,
                                'page': page_num+1,
                                'index': len(tables)+1
                            })
                            table_data = []
                            in_table = False
                    else:
                        # 简单判断是否为表格行（包含多个空格分隔的值）
                        if len(line.split()) > 2:
                            table_data.append(line.split())
                            in_table = True
            
            return tables
        except Exception as e:
            print(f"从PDF文件提取表格失败: {e}")
            return []
    
    def process_document(self, file_path):
        """处理文档中的多模态内容"""
        result = {
            'images': [],
            'tables': []
        }
        
        if file_path.endswith('.docx'):
            # 处理Word文件
            result['images'] = self.extract_images_from_word(file_path)
            result['tables'] = self.extract_tables_from_word(file_path)
        elif file_path.endswith('.pdf'):
            # 处理PDF文件
            result['images'] = self.extract_images_from_pdf(file_path)
            result['tables'] = self.extract_tables_from_pdf(file_path)
        
        return result
    
    def get_image_base64(self, image_path):
        """将图片转换为Base64编码"""
        try:
            with open(image_path, "rb") as f:
                image_bytes = f.read()
            return base64.b64encode(image_bytes).decode('utf-8')
        except Exception as e:
            print(f"转换图片为Base64失败: {e}")
            return None
